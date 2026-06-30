import re
import os
import sys
import datetime
import pdfplumber
import docx
from typing import Dict, Any, List, Optional, Tuple
from sources.notes_parser import COMMON_SKILLS

# Section header regexes (only match when they appear on their own line with optional trailing punctuation/spaces)
SKILLS_HEADERS = re.compile(r'^\s*(?:skills|technical skills|technologies|expertise|core competencies|skills & expertise|skills and tools|key skills|professional skills)\b[:\s]*$', re.IGNORECASE | re.MULTILINE)
EXPERIENCE_HEADERS = re.compile(r'^\s*(?:experience|work experience|employment history|work history|professional experience|employment|career history|work background)\b[:\s]*$', re.IGNORECASE | re.MULTILINE)
EDUCATION_HEADERS = re.compile(r'^\s*(?:education|academic background|academic history|education & credentials|academic qualifications)\b[:\s]*$', re.IGNORECASE | re.MULTILINE)
PROJECTS_HEADERS = re.compile(r'^\s*(?:projects|academic projects|personal projects|key projects|technical projects)\b[:\s]*$', re.IGNORECASE | re.MULTILINE)

# Contact info regexes
EMAIL_REGEX = re.compile(r'\b[a-zA-Z0-9._%+-]+\s*@\s*[a-zA-Z0-9.-]+\s*\.\s*[a-zA-Z]{2,}\b')
PHONE_REGEX = re.compile(r'\+?\d{1,4}[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}')
DATE_RANGE_REGEX = re.compile(r'\b((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s-]*\d{4}|\d{4}-\d{2}|\d{2}/\d{4}|\d{4})[-.\s]+(?:to|[-—])[-.\s]+((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s-]*\d{4}|\d{4}-\d{2}|\d{2}/\d{4}|\d{4}|Present|Current)\b', re.IGNORECASE)

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {str(e)}")
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    text = []
    try:
        # Wrap docx Document call inside a with open statement to explicitly release the file handle
        with open(file_path, "rb") as f:
            doc = docx.Document(f)
            for para in doc.paragraphs:
                text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {str(e)}")
    return "\n".join(text).strip()

def split_sections(text: str) -> Dict[str, str]:
    sections = {"header": text, "skills": "", "experience": "", "education": "", "projects": ""}
    
    matches = []
    for match in SKILLS_HEADERS.finditer(text):
        matches.append((match.start(), match.end(), "skills"))
    for match in EXPERIENCE_HEADERS.finditer(text):
        matches.append((match.start(), match.end(), "experience"))
    for match in EDUCATION_HEADERS.finditer(text):
        matches.append((match.start(), match.end(), "education"))
    for match in PROJECTS_HEADERS.finditer(text):
        matches.append((match.start(), match.end(), "projects"))
        
    if not matches:
        return sections
        
    # Sort matches by start index
    matches.sort(key=lambda x: x[0])
    
    # Slice the text
    sections["header"] = text[:matches[0][0]]
    
    for i in range(len(matches)):
        start_idx = matches[i][1]
        end_idx = matches[i+1][0] if i + 1 < len(matches) else len(text)
        sec_type = matches[i][2]
        # In case a section type appears multiple times, append it
        sections[sec_type] = (sections.get(sec_type, "") + "\n" + text[start_idx:end_idx]).strip()
        
    return sections

def parse_job_date_range(line: str) -> Optional[Tuple[str, str, int, int]]:
    dash_pat = r'(?:to|[-—–\u2013\u2014])'
    months = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*'
    
    pat1 = r'\b((?:' + months + r'[\s-]*\d{4}|\d{4}-\d{2}|\d{2}/\d{4}|\d{4}))[-.\s\u2013\u2014—–]+' + dash_pat + r'[-.\s\u2013\u2014—–]+((?:' + months + r'[\s-]*\d{4}|\d{4}-\d{2}|\d{2}/\d{4}|\d{4}|Present|Current))\b'
    pat2 = r'\b(' + months + r')[-.\s\u2013\u2014—–]+' + dash_pat + r'[-.\s\u2013\u2014—–]+((' + months + r')[\s-]*(\d{4}))\b'
    
    m1 = re.search(pat1, line, re.IGNORECASE)
    if m1:
        return m1.group(1).strip(), m1.group(2).strip(), m1.start(), m1.end()
        
    m2 = re.search(pat2, line, re.IGNORECASE)
    if m2:
        start_month = m2.group(1).strip()
        end_date = m2.group(2).strip()
        year = m2.group(4).strip()
        start_date = f"{start_month} {year}"
        return start_date, end_date, m2.start(), m2.end()
        
    return None

def parse_job_header(line: str) -> Optional[Tuple[str, str, str, str]]:
    date_res = parse_job_date_range(line)
    if not date_res:
        return None
        
    start_date, end_date, match_start, match_end = date_res
    header_part = line[:match_start].strip().rstrip("() ,-–—\u2013\u2014")
    
    KNOWN_COMPANIES = {
        "samsung", "google", "infosys", "wipro", "tcs", "cognizant", "accenture", "microsoft", 
        "amazon", "facebook", "meta", "netflix", "apple", "corporation", "corp", "ltd", "inc", 
        "co", "program", "solutions", "labs", "company", "institutes", "iihr", "acme", "springboard"
    }
    
    KNOWN_LOCATIONS = {
        "bangalore", "bengaluru", "gwalior", "jhansi", "nagpur", "mumbai", "pune", "chennai", 
        "hyderabad", "delhi", "kolkata", "noida", "gurgaon", "karnataka", "madhya pradesh", 
        "uttar pradesh", "maharashtra", "tamil nadu", "telangana", "andhra pradesh", "gujarat", 
        "rajasthan", "west bengal", "kerala", "punjab", "haryana", "bihar", "odisha", "india", "remote"
    }
    
    splitters = [r'\s*\|\s*', r'\s*(?:[\u2014\u2013—–]|\s-\s)\s*', r'\s*,\s*']
    
    parts = [header_part]
    for splitter in splitters:
        new_parts = []
        for p in parts:
            new_parts.extend(re.split(splitter, p))
        parts = [pt.strip() for pt in new_parts if pt.strip()]
        
    if len(parts) >= 2:
        company_idx = -1
        for idx, pt in enumerate(parts):
            pt_lower = pt.lower()
            if any(re.search(rf'\b{re.escape(kc)}\b', pt_lower) for kc in KNOWN_COMPANIES):
                company_idx = idx
                break
        
        if company_idx != -1:
            company = parts[company_idx]
            title_parts = []
            for idx, pt in enumerate(parts):
                if idx != company_idx:
                    pt_words = set(pt.lower().split())
                    if not pt_words.intersection(KNOWN_LOCATIONS):
                        title_parts.append(pt)
            title = " - ".join(title_parts) if title_parts else "Unknown"
        else:
            last_part = parts[-1]
            last_words = set(last_part.lower().split())
            if last_words.intersection(KNOWN_LOCATIONS) and len(parts) >= 3:
                company = parts[-2]
                title_parts = [parts[i] for i in range(len(parts) - 2)]
            else:
                company = last_part
                title_parts = parts[:-1]
                
            filtered_title_parts = []
            for pt in title_parts:
                pt_words = set(pt.lower().split())
                if not pt_words.intersection(KNOWN_LOCATIONS):
                    filtered_title_parts.append(pt)
            title = " - ".join(filtered_title_parts) if filtered_title_parts else "Unknown"
            
        return title, company, start_date, end_date
        
    elif len(parts) == 1:
        return parts[0], "Unknown", start_date, end_date
        
    return None

BOUNDARY_KEYWORDS = [
    "bengaluru", "bangalore", "gwalior", "jhansi", "nagpur", "mumbai", "pune", "chennai", "hyderabad", 
    "delhi", "kolkata", "noida", "gurgaon", "karnataka", "madhya pradesh", "uttar pradesh", 
    "maharashtra", "tamil nadu", "telangana", "andhra pradesh", "gujarat", "rajasthan", 
    "west bengal", "kerala", "punjab", "haryana", "bihar", "odisha", "india"
]

def find_city_in_text(text: str) -> Optional[str]:
    cities = ["bengaluru", "bangalore", "gwalior", "jhansi", "nagpur", "mumbai", "pune", "chennai", "hyderabad", "delhi", "kolkata", "noida", "gurgaon"]
    for city in cities:
        if re.search(r'\b' + re.escape(city) + r'\b', text, re.IGNORECASE):
            return city.title()
    return None

def is_institution_line(line: str) -> bool:
    inst_keywords = ["university", "college", "institute", "school", "academy", "iit", "iisc", "vnit", "bmsce", "sanskar public", "sun international"]
    return any(k in line.lower() for k in inst_keywords)

def is_degree_line(line: str) -> bool:
    deg_keywords = [
        "bachelor", "master", "doctor", "b.e.", "b.tech", "m.tech", "ph.d.", "b.s.", "m.s.", "b.a.",
        "hsc", "ssc", "cbse", "icse", "higher secondary", "central board", "secondary school", "certificate", "diploma"
    ]
    return any(k in line.lower() for k in deg_keywords)

def extract_institution_name(line: str) -> str:
    # Look for boundary keywords
    first_idx = len(line)
    matched_word = None
    
    for keyword in BOUNDARY_KEYWORDS:
        m = re.search(rf'\b{re.escape(keyword)}\b', line, re.IGNORECASE)
        if m:
            if m.start() < first_idx:
                first_idx = m.start()
                matched_word = keyword
                
    if matched_word:
        inst = line[:first_idx].strip()
        inst = inst.rstrip(",-–—\u2013\u2014 ")
        return inst
        
    parts = [p.strip() for p in line.split(",") if p.strip()]
    if parts:
        return parts[0]
    return line.strip()

def is_company_line(line: str) -> bool:
    """Checks if a line represents a company and location name without a date."""
    dash_pat = r'(?:to|[-—–\u2013\u2014])'
    date_pat = r'\b((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s-]*\d{4}|\d{4}-\d{2}|\d{2}/\d{4}|\d{4})[-.\s\u2013\u2014—–]+' + dash_pat + r'[-.\s\u2013\u2014—–]+((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s-]*\d{4}|\d{4}-\d{2}|\d{2}/\d{4}|\d{4}|Present|Current)\b'
    if re.search(date_pat, line, re.IGNORECASE):
        return False
    cities = ["bengaluru", "bangalore", "gwalior", "jhansi", "nagpur", "mumbai", "pune", "chennai", "hyderabad", "delhi", "kolkata", "noida", "gurgaon", "india"]
    company_keywords = ["institute", "research", "university", "college", "solutions", "labs", "inc", "corp", "co", "ltd", "technologies", "company"]
    has_city = any(c in line.lower() for c in cities)
    has_keyword = any(k in line.lower() for k in company_keywords)
    return has_city or has_keyword

def extract_hyperlinks_from_pdf(file_path: str) -> List[str]:
    """Extracts embedded hyperlinks from a PDF file using pdfplumber."""
    links = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                hyperlinks = page.hyperlinks
                if hyperlinks:
                    for hl in hyperlinks:
                        uri = hl.get("uri")
                        if uri:
                            uri_str = uri.strip()
                            if uri_str.lower().startswith("http://") or uri_str.lower().startswith("https://"):
                                links.append(uri_str)
    except Exception as e:
        print(f"Error reading hyperlinks from PDF {file_path}: {str(e)}")
    return links

def parse_experience_section(exp_text: str) -> List[Dict[str, Any]]:
    experience_entries = []
    if not exp_text:
        return experience_entries

    stop_keywords = {
        "projects", "certifications", "skills summary", "education", 
        "extracurricular activities", "skills", "technical skills", 
        "interests", "publications", "awards", "courses"
    }

    raw_lines = exp_text.split("\n")
    lines = []
    for line in raw_lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        if line_clean.lower() in stop_keywords:
            break
        lines.append(line_clean)
        
    current_job = None
    has_valid_header = False

    i = 0
    while i < len(lines):
        is_two_line_job = False
        if i + 1 < len(lines):
            line1 = lines[i]
            line2 = lines[i+1]
            if is_company_line(line1) and parse_job_header(line2):
                is_two_line_job = True
                
        if is_two_line_job:
            line1 = lines[i]
            line2 = lines[i+1]
            
            company = extract_institution_name(line1)
            date_res = parse_job_date_range(line2)
            start_date, end_date, match_start, match_end = date_res
            title = line2[:match_start].strip().rstrip("() ,-–—\u2013\u2014")
            
            current_job = {
                "company": company.strip(),
                "title": title.strip(),
                "start": start_date,
                "end": end_date,
                "summary": ""
            }
            experience_entries.append(current_job)
            has_valid_header = True
            i += 2
        else:
            line = lines[i]
            header_info = parse_job_header(line)
            
            if header_info:
                title, company_loc, start_date, end_date = header_info
                
                comp_parts = [p.strip() for p in company_loc.split(",") if p.strip()]
                company = comp_parts[0] if comp_parts else "Unknown"
                
                current_job = {
                    "company": company.strip(),
                    "title": title.strip(),
                    "start": start_date,
                    "end": end_date,
                    "summary": ""
                }
                experience_entries.append(current_job)
                has_valid_header = True
            else:
                line_clean = re.sub(r'^[-*•\u2022]\s*', '', line).strip()
                if current_job is not None:
                    if current_job["summary"]:
                        current_job["summary"] += "\n" + line_clean
                    else:
                        current_job["summary"] = line_clean
            i += 1

    for entry in experience_entries:
        if not entry["summary"]:
            entry["summary"] = None

    if not has_valid_header:
        print("[WARNING] No valid experience headers found in experience section. Returning empty list.", file=sys.stderr)
        return []

    return experience_entries

def is_section_header_line(line: str) -> bool:
    line_clean = line.strip().upper().rstrip(":-—–•* ")
    keywords = [
        "SKILLS", "TECHNICAL SKILLS", "SKILLS & TECHNOLOGIES", "SKILLS AND TECHNOLOGIES",
        "EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY", "PROFESSIONAL EXPERIENCE",
        "PROJECTS", "ACADEMIC PROJECTS", "PERSONAL PROJECTS",
        "ACHIEVEMENTS", "AWARDS", "CERTIFICATIONS", "LICENSES & CERTIFICATIONS", "COURSES",
        "PUBLICATIONS", "INTERESTS", "EXTRACURRICULAR ACTIVITIES", "LANGUAGES", "CORE SUBJECTS", 
        "RELEVANT COURSEWORK", "COURSEWORK", "DEVELOPER TOOLS", "DATABASES", "PROGRAMMING LANGUAGES"
    ]
    if line_clean in keywords:
        return True
    for k in keywords:
        if line_clean.startswith(k) and (len(line_clean) == len(k) or line_clean[len(k)] in [" ", "&", "/", "|", ":", "-", "–", "—"]):
            if len(line_clean.split()) <= 4:
                return True
    return False

def parse_education_block(block_lines: List[str]) -> Dict[str, Any]:
    DEGREE_REGEX = re.compile(
        r'\b(B\.?E\b\.?|B\.?Tech\b\.?|M\.?Tech\b\.?|B\.?S\b\.?|M\.?S\b\.?|Ph\.?D\b\.?|Bachelor\s+of\s+Science\b|Bachelor\s+of\s+Engineering\b|Bachelor\s+of\s+Technology\b|Bachelor\s+of\s+Arts\b|Master\s+of\s+Science\b|Master\s+of\s+Engineering\b|Master\s+of\s+Technology\b|Master\s+of\s+Business\s+Administration\b|Bachelor[s]?\b|Master[s]?\b|Doctor\b|MBA\b|B\.?A\b\.?|Higher\s+Secondary\s+Certificate\s*\(HSC\)|Higher\s+Secondary\s+Certificate\b|Higher\s+Secondary\b|Secondary\s+School\s+Certificate\s*\(SSC\)|Secondary\s+School\s+Certificate\b|Secondary\s+School\b|Central\s+Board\s+of\s+Secondary\s+Education\b|Central\s+Board\b|Indian\s+Certificate\s+of\s+Secondary\s+Education\b|Certificate\b|Diploma\b|HSC\b|SSC\b|CBSE\b|ICSE\b|Class\s+XII\b|Class\s+X\b)', 
        re.IGNORECASE
    )
    
    # 1. Find year
    years = []
    for line in block_lines:
        found_years = re.findall(r'\b(20\d{2}|19\d{2})\b', line)
        years.extend([int(y) for y in found_years])
    end_year = years[-1] if years else None
    
    # 2. Find institution
    institution = None
    for line in block_lines:
        parts = re.split(r'\s*(?:,|\b-\b|\b–\b|\b—\b|\||·)\s*', line)
        for part in parts:
            part_clean = part.strip()
            if is_institution_line(part_clean) and not DEGREE_REGEX.search(part_clean):
                institution = extract_institution_name(part_clean)
                break
        if institution:
            break
            
    if not institution:
        for line in block_lines:
            parts = re.split(r'\s*(?:,|\b-\b|\b–\b|\b—\b|\||·)\s*', line)
            for part in parts:
                part_clean = part.strip()
                if is_institution_line(part_clean):
                    institution = extract_institution_name(part_clean)
                    break
            if institution:
                break
                
    if not institution and block_lines:
        first_line = block_lines[0]
        if not any(k in first_line.lower() for k in ["gpa", "cgpa", "grade", "percentage", "marks", "class "]) and not DEGREE_REGEX.search(first_line):
            institution = extract_institution_name(first_line)
        else:
            for line in block_lines[1:]:
                if not DEGREE_REGEX.search(line) and not any(k in line.lower() for k in ["gpa", "cgpa", "grade", "percentage", "marks", "%"]):
                    institution = extract_institution_name(line)
                    break
            if not institution:
                institution = "Unknown"
                
    # 3. Find degree and field
    degree = None
    field = None
    
    for line in block_lines:
        line_clean = re.sub(r'\(\s*\d{4}\s*[-–—to\s]+\s*(?:\d{4}|Present|Current)\s*\)', '', line, flags=re.IGNORECASE)
        line_clean = re.sub(r'\b\d{4}\s*[-–—to\s]+\s*(?:\d{4}|Present|Current)\b', '', line_clean, flags=re.IGNORECASE)
        line_clean = re.sub(r'\(\s*\d{4}\s*\)', '', line_clean)
        line_clean = re.sub(r'\b\d{4}\b', '', line_clean)
        line_clean = re.sub(r'\(\s*\)', '', line_clean)
        line_clean = line_clean.strip()
        
        deg_match = DEGREE_REGEX.search(line_clean)
        if deg_match:
            degree = deg_match.group(1).strip()
            # Split line to find field
            parts = re.split(r'\s*(?:[\u2014\u2013—–\uFFFD]| \- | \| |,|:|·|;)\s*', line_clean)
            field_parts = []
            for part in parts:
                part_clean = part.strip()
                if not part_clean:
                    continue
                # Skip if it is the matched degree itself (exact match)
                part_deg_clean = re.sub(r'[^a-zA-Z0-9]', '', part_clean).lower()
                deg_clean_val = re.sub(r'[^a-zA-Z0-9]', '', degree).lower()
                if part_deg_clean == deg_clean_val:
                    continue
                if is_institution_line(part_clean):
                    continue
                # Skip numeric grades, GPAs, or percentage values
                if re.match(r'^\s*[\d\.\s%/-]+\s*$', part_clean) or any(k in part_clean.lower() for k in ["gpa", "cgpa", "grade", "percentage", "marks", "%", "/"]):
                    continue
                if any(k in part_clean.lower() for k in BOUNDARY_KEYWORDS):
                    continue
                # Strip degree keyword out of the field part if it's there
                deg_word = degree.rstrip(".")
                part_clean = re.sub(rf'\b{re.escape(deg_word)}\b\.?', '', part_clean, flags=re.IGNORECASE).strip()
                # Strip years from the part
                part_clean = re.sub(r'\b(20\d{2}|19\d{2})\b', '', part_clean).strip()
                if part_clean:
                    field_parts.append(part_clean)
            if field_parts:
                field = " ".join(field_parts)
            break
            
    if degree and not field:
        for line in block_lines:
            line_clean = re.sub(r'\b\d{4}\b', '', line).strip()
            if not line_clean or is_institution_line(line_clean) or DEGREE_REGEX.search(line_clean):
                continue
            if any(k in line_clean.lower() for k in ["gpa", "cgpa", "grade", "percentage", "marks", "%", "/"]):
                continue
            field = line_clean
            break
            
    if field:
        field = re.sub(r'^(?:in|of|major in)\s+', '', field, flags=re.IGNORECASE).strip()
        field = field.strip("-*• \t,;|")
        if not field or len(field) > 60:
            field = None
            
    if degree:
        degree = degree.strip("-*• \t,;|")
        
    return {
        "institution": institution or "Unknown",
        "degree": degree,
        "field": field,
        "end_year": end_year
    }

def parse_education_section(edu_text: str) -> List[Dict[str, Any]]:
    education_entries = []
    if not edu_text:
        return education_entries

    raw_lines = edu_text.split("\n")
    lines = []
    for line in raw_lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        # Stop immediately when any section boundary is hit
        if is_section_header_line(line_clean):
            break
        lines.append(line_clean)
        
    # Group lines into blocks representing individual education records
    blocks = []
    current_block = []
    for line in lines:
        starts_new = False
        if not current_block:
            starts_new = True
        else:
            # Institution lines start a new block
            if is_institution_line(line):
                starts_new = True
            # Year on a line starts a new block if the current block already contains a year
            elif re.search(r'\b(20\d{2}|19\d{2})\b', line):
                has_year = any(re.search(r'\b(20\d{2}|19\d{2})\b', l) for l in current_block)
                if has_year:
                    starts_new = True
                    
        if starts_new and current_block:
            blocks.append(current_block)
            current_block = []
        current_block.append(line)
        
    if current_block:
        blocks.append(current_block)
        
    # Parse each block
    for block in blocks:
        entry = parse_education_block(block)
        education_entries.append(entry)
        
    return education_entries

def parse_location_from_header(header_text: str) -> Dict[str, Optional[str]]:
    loc = {"city": None, "region": None, "country": None}
    
    known_cities = [
        "bangalore", "bengaluru", "mumbai", "pune", "delhi", "new delhi", "chennai", 
        "hyderabad", "kolkata", "gurgaon", "gurugram", "noida", "ahmedabad", "jaipur",
        "gwalior", "jhansi", "nagpur", "chittoor", "tirupati", "gangavathi"
    ]
    
    known_regions = [
        "karnataka", "maharashtra", "delhi", "tamil nadu", "telangana", "west bengal", 
        "haryana", "uttar pradesh", "gujarat", "rajasthan", "madhya pradesh", "andhra pradesh"
    ]
    
    lines = [line.strip() for line in header_text.split("\n") if line.strip()]
    for line in lines:
        city_match = None
        for city in known_cities:
            m = re.search(rf'\b{re.escape(city)}\b', line, re.IGNORECASE)
            if m:
                city_match = city.capitalize()
                break
                
        region_match = None
        for region in known_regions:
            m = re.search(rf'\b{re.escape(region)}\b', line, re.IGNORECASE)
            if m:
                region_match = region.title()
                break
                
        country_match = None
        if re.search(r'\b(India|IN|USA|US)\b', line, re.IGNORECASE):
            country_match = "IN" if re.search(r'\b(India|IN)\b', line, re.IGNORECASE) else "US"
            
        if city_match:
            loc["city"] = city_match
            if region_match:
                loc["region"] = region_match
            if country_match:
                loc["country"] = country_match
            return loc
            
    return loc

def parse_location(text: str) -> Dict[str, Optional[str]]:
    """
    Parses location from text.
    Looks for line starting with "Location:" followed by city, state, country.
    Falls back to a header-wide scan if "Location:" label is missing.
    """
    loc = {"city": None, "region": None, "country": None}
    match = re.search(r'\bLocation\s*:\s*([^\n]+)', text, re.IGNORECASE)
    if match:
        loc_str = match.group(1).strip()
        parts = [p.strip() for p in loc_str.split(",") if p.strip()]
        if len(parts) == 1:
            loc["city"] = parts[0]
        elif len(parts) == 2:
            loc["city"] = parts[0]
            val = parts[1]
            if val.lower() in ["india", "ind", "usa", "united states", "us", "uk", "united kingdom", "canada", "ca", "germany", "de"]:
                loc["country"] = val
            else:
                loc["region"] = val
        elif len(parts) >= 3:
            loc["city"] = parts[0]
            loc["region"] = parts[1]
            loc["country"] = parts[2]
            
    if not loc["city"]:
        sections = split_sections(text)
        header_text = sections.get("header", text[:1000])
        header_loc = parse_location_from_header(header_text)
        if header_loc["city"]:
            loc = header_loc
            
    return loc

def parse_year_month(date_str: str) -> Optional[Tuple[int, int]]:
    """Helper to parse (year, month) from various date formats."""
    if not date_str:
        return None
    # Format 1: 2023-06
    m = re.search(r'\b(20\d{2}|19\d{2})-(\d{2})\b', date_str)
    if m:
        return int(m.group(1)), int(m.group(2))
    
    # Format 2: Jan 2023
    months_map = {
        "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
        "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12
    }
    m_name = re.search(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s-]*\b(20\d{2}|19\d{2})\b', date_str, re.IGNORECASE)
    if m_name:
        mon_str = m_name.group(1).lower()[:3]
        return int(m_name.group(2)), months_map.get(mon_str, 1)
        
    # Format 3: 2023 (just year)
    m_yr = re.search(r'\b(20\d{2}|19\d{2})\b', date_str)
    if m_yr:
        return int(m_yr.group(1)), 1
        
    return None

def is_skill_heading(s: str) -> bool:
    s_lower = s.lower().strip()
    headings = [
        "languages", "backend", "frontend", "tools", "core subjects", "coursework", 
        "skills", "technical skills", "frameworks", "programming languages", "databases", 
        "devops", "ai / ml", "cs fundamentals", "technologies", "development", "subjects",
        "core", "relevant coursework", "programming"
    ]
    for h in headings:
        if h in s_lower:
            # Check if it has a comma or is longer, i.e., it's a list, not a heading
            if "," not in s_lower and len(s_lower.split()) <= 4:
                return True
    return False

def clean_headline_text(raw_val: str) -> Optional[str]:
    if not raw_val:
        return None
    lines = raw_val.split("\n")
    cleaned_lines = []
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        if is_section_header_line(line_clean):
            break
        cleaned_lines.append(line_clean)
    if not cleaned_lines:
        return None
    val = " ".join(cleaned_lines)
    val = re.sub(r'^\s*[-*•\u2022]\s*', '', val)
    return " ".join(val.split()[:20]).strip()

def is_tech_stack_line(line: str) -> bool:
    line_clean = line.strip().rstrip(".-*• \t")
    if not line_clean:
        return False
    if line_clean.startswith(("-", "–", "—", "•", "*")):
        return False
    tech_keywords = [
        "python", "java", "c++", "c#", "javascript", "typescript", "react", "node", "express", 
        "mongodb", "sql", "mysql", "postgresql", "html", "css", "aws", "gcp", "docker", "kubernetes", 
        "git", "machine learning", "nlp", "deep learning", "opencv", "scikit", "streamlit", "flask", 
        "django", "spring", "bootstrap", "tailwind", "jquery", "pytorch", "tensorflow", "keras", "jwt", "file handling"
    ]
    parts = [p.strip().lower() for p in re.split(r'[,|·]', line_clean) if p.strip()]
    if not parts:
        return False
    matched = 0
    for part in parts:
        if any(tk in part for tk in tech_keywords) or len(part.split()) <= 2:
            matched += 1
    return matched / len(parts) >= 0.7 and len(parts) >= 2

def split_title_and_tech(line: str) -> Tuple[str, Optional[str]]:
    tech_keywords = {
        "python", "scikit-learn", "streamlit", "react", "react.js", "node", "node.js", 
        "express", "express.js", "mongodb", "jwt", "html", "css", "javascript", "c++", 
        "java", "sql", "c", "mysql", "postgresql", "aws", "gcp", "docker", "kubernetes", 
        "git", "ml", "nlp", "opencv", "flask", "django", "spring"
    }
    words = line.split()
    tech_words = []
    title_words = []
    in_tech = True
    for word in reversed(words):
        word_clean = word.strip(",;|()").lower()
        if in_tech and (word_clean in tech_keywords or word_clean in [",", "&", "and"]):
            tech_words.insert(0, word)
        else:
            in_tech = False
            title_words.insert(0, word)
            
    if tech_words and title_words:
        title = " ".join(title_words).strip().rstrip(",-–—\u2013\u2014 ")
        tech = " ".join(tech_words).strip().strip(",-–—\u2013\u2014 ")
        return title, tech
    return line, None

def extract_projects(proj_text: str) -> List[Dict[str, Any]]:
    projects = []
    if not proj_text:
        return projects

    raw_lines = proj_text.split("\n")
    lines = []
    for line in raw_lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        if is_section_header_line(line_clean):
            break
        lines.append(line_clean)
        
    def is_project_header(line: str, is_first: bool) -> bool:
        if is_first:
            return True
        if "|" in line:
            return True
        if re.search(r'\b(Ongoing|Present|Current|20\d{2}|19\d{2})\b', line, re.IGNORECASE):
            return True
        title, tech = split_title_and_tech(line)
        if tech:
            return True
        proj_nouns = ["chatbot", "tracker", "platform", "splitease", "app", "system", "tool", "website", "splitter"]
        if any(noun in line.lower() for noun in proj_nouns):
            return True
        return False

    grouped = []
    current_proj = None
    for line in lines:
        line_clean = line.strip()
        
        is_bullet = False
        for bullet in ["•", "–", "—", "*", "-"]:
            if line_clean.startswith(bullet):
                is_bullet = True
                line_clean = line_clean[len(bullet):].strip()
                break
                
        if is_bullet:
            if current_proj:
                current_proj["description_lines"].append(line_clean)
        else:
            if is_project_header(line_clean, len(grouped) == 0 and current_proj is None):
                if current_proj:
                    grouped.append(current_proj)
                current_proj = {
                    "title_line": line_clean,
                    "tech_stack_raw": None,
                    "description_lines": []
                }
            elif current_proj and not current_proj["tech_stack_raw"] and is_tech_stack_line(line_clean):
                current_proj["tech_stack_raw"] = line_clean
            else:
                if current_proj:
                    current_proj["description_lines"].append(line_clean)
                else:
                    current_proj = {
                        "title_line": line_clean,
                        "tech_stack_raw": None,
                        "description_lines": []
                    }
                
    if current_proj:
        grouped.append(current_proj)
        
    parsed_projects = []
    for proj in grouped:
        title_line = proj["title_line"]
        tech_stack_raw = proj["tech_stack_raw"]
        desc_lines = proj["description_lines"]
        
        duration = None
        duration_match = re.search(r'[-—–\u2013\u2014|]\s*(Ongoing|Present|Current|20\d{2}|19\d{2})\b', title_line, re.IGNORECASE)
        if duration_match:
            duration = duration_match.group(1).strip()
            title_line = title_line[:duration_match.start()].strip()
        else:
            duration_match = re.search(r'\b(Ongoing|Present|Current|20\d{2}|19\d{2})\b', title_line, re.IGNORECASE)
            if duration_match:
                duration = duration_match.group(1).strip()
                title_line = title_line[:duration_match.start()].strip().rstrip(" -—–\u2013\u2014(|")
                
        title = title_line
        tech_stack_list = []
        
        if "|" in title_line:
            parts = [p.strip() for p in title_line.split("|") if p.strip()]
            if parts:
                title = parts[0]
                if len(parts) >= 2:
                    possible_tech = parts[1]
                    if not duration:
                        dur_m = re.search(r'\b(Ongoing|Present|Current|20\d{2}|19\d{2})\b', possible_tech, re.IGNORECASE)
                        if dur_m:
                            duration = dur_m.group(1).strip()
                            possible_tech = possible_tech[:dur_m.start()].strip().rstrip(" -—–\u2013\u2014(|")
                    tech_stack_list = [t.strip() for t in re.split(r'[,;]', possible_tech) if t.strip()]
        else:
            title, tech_stack_str = split_title_and_tech(title_line)
            if tech_stack_str:
                tech_stack_list = [t.strip() for t in re.split(r'[,;]', tech_stack_str) if t.strip()]
                
        if not tech_stack_list and tech_stack_raw:
            tech_stack_list = [t.strip() for t in re.split(r'[,;]', tech_stack_raw) if t.strip()]
            
        description = "\n".join(desc_lines).strip() if desc_lines else None
        title = title.strip(".-*• \t,;|—–\u2013\u2014")
        
        parsed_projects.append({
            "name": title,
            "description": description,
            "tech_stack": tech_stack_list,
            "duration": duration,
            "source": "resume"
        })
        
    return parsed_projects

def parse_resume(file_path: str) -> Dict[str, Any]:
    """
    Parses resume PDF or DOCX file.
    Extracts text, runs regex for email, phone, and splits into sections.
    Handles scanned or empty resumes gracefully by returning null fields.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        text = extract_text_from_docx(file_path)
    else:
        text = ""

    candidate = {
        "full_name": None,
        "emails": [],
        "phones": [],
        "location": {"city": None, "region": None, "country": None},
        "links": {"linkedin": None, "github": None, "leetcode": None, "hackerrank": None, "portfolio": None, "other": []},
        "headline": None,
        "years_experience": None,
        "skills": [],
        "experience": [],
        "education": [],
        "projects": []
    }

    if not text:
        print(f"Resume text is empty (scanned PDF or empty file): {file_path}")
        return candidate

    # Extract emails and phones from the entire text
    emails = []
    for match in EMAIL_REGEX.findall(text):
        email_clean = re.sub(r'\s+', '', match).lower()
        emails.append(email_clean)
    emails = sorted(list(set(emails)))
    
    phones = list(set(PHONE_REGEX.findall(text)))
    candidate["emails"] = emails
    candidate["phones"] = phones

    # Extract location from the entire text
    candidate["location"] = parse_location(text)

    # Extract LinkedIn, GitHub, LeetCode, HackerRank and Portfolio/Website links from the entire text using regex
    # Regex to find links (http, https, or bare domains/profiles)
    url_pattern = r'\b(?:https?://)?(?:www\.)?(?:linkedin\.com/in/[a-zA-Z0-9_/-]+|github\.com/[a-zA-Z0-9_/-]+|leetcode\.com/[a-zA-Z0-9_/-]+|hackerrank\.com/[a-zA-Z0-9_/-]+|[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}/[^\s,;]*)'
    raw_urls = re.findall(url_pattern, text, re.IGNORECASE)
    
    hyperlinks = []
    # Also merge URLs from PDF hyperlinks if available
    if ext == ".pdf":
        try:
            hyperlinks = [h.strip() for h in extract_hyperlinks_from_pdf(file_path) if h.strip()]
            raw_urls.extend(hyperlinks)
        except Exception:
            pass
            
    # Clean and filter collected URLs
    linkedin_urls = []
    github_urls = []
    leetcode_urls = []
    hackerrank_urls = []
    portfolio_urls = []
    
    for url in raw_urls:
        url_clean = url.strip().rstrip("/.")
        if not url_clean:
            continue
        # Ensure it has a protocol if it's a link
        if not (url_clean.startswith("http://") or url_clean.startswith("https://")):
            url_clean_full = f"https://{url_clean}"
        else:
            url_clean_full = url_clean
            
        url_lower = url_clean_full.lower()
        if "mailto:" in url_lower:
            continue
            
        if "linkedin.com/in/" in url_lower:
            if url_clean_full not in linkedin_urls:
                linkedin_urls.append(url_clean_full)
        elif "github.com/" in url_lower:
            # Avoid blob, tree, etc.
            if not any(p in url_lower for p in ["/blob/", "/tree/", "/issues", "/pull"]):
                if url_clean_full not in github_urls:
                    github_urls.append(url_clean_full)
        elif "leetcode.com/" in url_lower:
            if url_clean_full not in leetcode_urls:
                leetcode_urls.append(url_clean_full)
        elif "hackerrank.com/" in url_lower:
            if url_clean_full not in hackerrank_urls:
                hackerrank_urls.append(url_clean_full)
        else:
            # Potential portfolio website
            # Filter out generic or false positive patterns
            if any(domain in url_lower for domain in ["github.com", "linkedin.com", "leetcode.com", "hackerrank.com"]):
                continue
            if url_clean_full not in portfolio_urls:
                portfolio_urls.append(url_clean_full)
                
    # Populate LinkedIn
    if linkedin_urls:
        candidate["links"]["linkedin"] = linkedin_urls[0]
        if any(h.lower() in linkedin_urls[0].lower() for h in hyperlinks):
            candidate["_hyperlink_linkedin"] = True
        
    # Populate GitHub
    if github_urls:
        candidate["links"]["github"] = github_urls[0]
        if any(h.lower() in github_urls[0].lower() for h in hyperlinks):
            candidate["_hyperlink_github"] = True
        
    # Populate LeetCode
    if leetcode_urls:
        candidate["links"]["leetcode"] = leetcode_urls[0]
        if any(h.lower() in leetcode_urls[0].lower() for h in hyperlinks):
            candidate["_hyperlink_leetcode"] = True

    # Populate HackerRank
    if hackerrank_urls:
        candidate["links"]["hackerrank"] = hackerrank_urls[0]
        if any(h.lower() in hackerrank_urls[0].lower() for h in hyperlinks):
            candidate["_hyperlink_hackerrank"] = True

    # Populate Portfolio website
    if portfolio_urls:
        candidate["links"]["portfolio"] = portfolio_urls[0]
        if any(h.lower() in portfolio_urls[0].lower() for h in hyperlinks):
            candidate["_hyperlink_portfolio"] = True
            
    # Put remaining portfolio websites or links in links.other if they are not chosen
    for p_url in portfolio_urls[1:] if (portfolio_urls and candidate["links"]["portfolio"] == portfolio_urls[0]) else portfolio_urls:
        if p_url not in candidate["links"]["other"]:
            candidate["links"]["other"].append(p_url)
            
    # Add extra leetcodes/hackerranks to links.other
    if len(leetcode_urls) > 1:
        for extra in leetcode_urls[1:]:
            if extra not in candidate["links"]["other"]:
                candidate["links"]["other"].append(extra)
    if len(hackerrank_urls) > 1:
        for extra in hackerrank_urls[1:]:
            if extra not in candidate["links"]["other"]:
                candidate["links"]["other"].append(extra)

    # Split text into sections
    sections = split_sections(text)

    # Extract name from header section
    header_lines = [line.strip() for line in sections["header"].split("\n") if line.strip()]
    name = None
    for line in header_lines:
        # Name shouldn't contain email, phone, or links
        if "@" in line or any(c in line for c in ["/", "http", "www", "github", "linkedin"]):
            continue
        # Name shouldn't be too long or short
        if 2 <= len(line.split()) <= 4:
            name = line
            break
    candidate["full_name"] = name
    
    # Extract headline from resume text
    headline = None
    
    # Priority 1: Professional Summary / Summary / About Me
    summary_match = re.search(r'\b(?:Professional Summary|Summary|About Me)\b[:\s]*([^\n]+(?:\n[^\n]+){0,2})', text, re.IGNORECASE)
    if summary_match:
        val = clean_headline_text(summary_match.group(1))
        if val and len(val) > 10 and not any(k in val.lower() for k in ["skills:", "languages:", "coursework:", "relevant coursework:"]):
            headline = val

    # Priority 2: Career Objective
    if not headline:
        objective_match = re.search(r'\b(?:Career Objective)\b[:\s]*([^\n]+(?:\n[^\n]+){0,2})', text, re.IGNORECASE)
        if objective_match:
            val = clean_headline_text(objective_match.group(1))
            if val and len(val) > 10 and not any(k in val.lower() for k in ["skills:", "languages:", "coursework:", "relevant coursework:"]):
                headline = val

    # Priority 3: Objective
    if not headline:
        objective_match = re.search(r'\b(?:Objective)\b[:\s]*([^\n]+(?:\n[^\n]+){0,2})', text, re.IGNORECASE)
        if objective_match:
            val = clean_headline_text(objective_match.group(1))
            if val and len(val) > 10 and not any(k in val.lower() for k in ["skills:", "languages:", "coursework:", "relevant coursework:"]):
                headline = val
                
    candidate["headline"] = headline

    # Parse skills
    SKILL_BLACKLIST = {
        "coursework", "languages", "tools", "systems", "mathematics", "discrete", "engineering", 
        "databasemanagement", "datastructures&algorithms", "software engineering", "web & frameworks", 
        "cryptography & network", "computer organization & architecture", "operating systems",
        "relevant coursework", "programming languages", "technical skills", "web technologies", 
        "databases", "other skills problem solving", "programming languages c", "other skills",
        "web development", "databases & tools", "developer tools", "computer science", "frameworks",
        "database management", "data structures", "algorithms", "problem solving", "networks",
        "computer networks", "cryptography", "network security", "computer organization",
        "architecture", "operating system", "data structures & algorithms", "web", "framework"
    }
    extracted_skills = []
    skills_text = sections["skills"]
    if skills_text:
        # Split skills by comma, semicolon, newline, or tab
        for s in re.split(r'[,;\n\t]', skills_text):
            if ":" in s:
                s = s.split(":", 1)[1]
            s_clean = s.strip().strip(".-*• \t")
            if s_clean and len(s_clean) < 40:
                if is_skill_heading(s_clean):
                    continue
                s_lower = s_clean.lower()
                if s_lower in SKILL_BLACKLIST or any(b == s_lower for b in SKILL_BLACKLIST):
                    continue
                extracted_skills.append(s_clean)

    # Fallback/Additional scan: scan the entire resume text for common skills
    for skill in COMMON_SKILLS:
        skill_lower = skill.lower()
        if skill_lower in SKILL_BLACKLIST or is_skill_heading(skill):
            continue
        if re.search(rf'\b{re.escape(skill)}\b', text, re.IGNORECASE):
            extracted_skills.append(skill)

    candidate["skills"] = list(set(extracted_skills))

    # Parse experience
    candidate["experience"] = parse_experience_section(sections["experience"])

    # Parse education
    candidate["education"] = parse_education_section(sections["education"])

    # Parse projects
    candidate["projects"] = extract_projects(sections["projects"])

    # Fallback location.city inference if not explicitly provided
    if not candidate["location"].get("city"):
        inferred_city = None
        if sections["experience"]:
            inferred_city = find_city_in_text("\n".join(sections["experience"].split("\n")[:3]))
        if not inferred_city and sections["education"]:
            inferred_city = find_city_in_text("\n".join(sections["education"].split("\n")[:3]))
            
        if inferred_city:
            candidate["location"]["city"] = inferred_city
            candidate["_inferred_location_city"] = True

    # Attempt to calculate years of experience from experience dates
    earliest_date = None
    earliest_str = None
    for exp in candidate["experience"]:
        start_str = exp.get("start")
        if start_str:
            ym = parse_year_month(start_str)
            if ym:
                if earliest_date is None or ym < earliest_date:
                    earliest_date = ym
                    earliest_str = start_str
                    
    if earliest_date:
        today = datetime.date.today()
        current_year = today.year
        current_month = today.month
        
        start_year, start_month = earliest_date
        diff_months = (current_year - start_year) * 12 + (current_month - start_month)
        years_exp = round(max(0.0, diff_months / 12.0), 1)
        
        candidate["years_experience"] = years_exp
        
        # Print debug calculation for verification
        print(f"[DEBUG] Years of Experience Calculation for {candidate['full_name'] or 'Unknown'}:", file=sys.stderr)
        print(f"[DEBUG]   Earliest Start Date: {earliest_str} ({start_year}-{start_month:02d})", file=sys.stderr)
        print(f"[DEBUG]   Current Date:        {today.strftime('%Y-%m-%d')} ({current_year}-{current_month:02d})", file=sys.stderr)
        print(f"[DEBUG]   Months Difference:   {diff_months} months", file=sys.stderr)
        print(f"[DEBUG]   Calculated Years:    {years_exp} years", file=sys.stderr)

    return candidate

def parse_resumes(file_paths: List[str]) -> List[Dict[str, Any]]:
    """
    Parses multiple resume files independently.
    Returns a list of parsed candidate dictionaries.
    """
    candidates = []
    for path in file_paths:
        try:
            candidates.append(parse_resume(path))
        except Exception as e:
            print(f"Error parsing resume {path}: {str(e)}")
    return candidates
